import os
from docx import Document

doc = Document()
doc.add_heading('Sprint 7 Report: Cloud Deployment & Presentation Polish', 0)

content = """
Project: MediWise Intelligent Diagnosis System  
Sprint Focus: Cloud Deployment (Vercel, Render, MongoDB Atlas) and Clinical UI Redesign.  
Status: Completed Successfully

---

## 1. Executive Summary
Sprint 7 finalized the clinical portal by migrating all services from a local development environment to a secure, production-ready cloud infrastructure. We successfully launched the React.js client on Vercel's global edge network, deployed the FastAPI backend on Render.com, and configured a cloud-based MongoDB Atlas database. Furthermore, we implemented a major visual overhaul of the UI, incorporating clinical metrics, animated EKG waveforms, specialist referral routing, and PDF report exports. The system is live, secure, and ready for deployment.

## 2. Technical Accomplishments

### A. Cloud Database Migration (MongoDB Atlas)
- Set up a highly available MongoDB Atlas M0 cluster instance.
- Migrated the application configuration from local host dependencies to read connection strings from secured environment variables.
- Configured a database user account and opened network firewall rules (0.0.0.0/0) to enable access from Render's cloud servers while keeping administrative access restricted.

### B. Python API Deployment (Render.com)
- Deployed the FastAPI server to Render as an active web service.
- Solved the 512MB RAM memory ceiling limitation of Render's free tier by replacing the heavy default tensorflow package with tensorflow-cpu in backend/requirements.txt.
- Configured PYTHON_VERSION=3.10.11 and PYTHONPATH=backend on Render to ensure module imports are resolved properly and dependencies compile correctly.

### C. Model Deserialization Fixes (Keras 3 Compatibility)
- Solved Keras version serialization mismatch errors (e.g. quantization_config and renorm attributes missing across environments).
- Rewrote the loading logic to dynamically reconstruct both the Symptom ANN and Skin Cancer CNN models in Python, and loaded raw weights directly from symptom_ann_weights.weights.h5 and skin_cancer_cnn_weights.weights.h5.

### D. Frontend Deployment (Vercel)
- Deployed the React Vite single-page application on Vercel.
- Created a root-level vercel.json configuration file to instruct Vercel to build strictly the frontend/ folder as a static Vite build, bypassing standard monorepo folder parsing errors.
- Wired the live Render API URL (https://mediwise-api.onrender.com) into Vercel's VITE_API_URL environment variable.

### E. Clinical Logic Integration & Naming
- Renamed the "Dermatology Scanner" to "Image Scanner" on the UI and code, generalising the feature for clinical diagnostics.
- Verified and mapped bleeding symptoms to target prognosis logic: Nose Bleeding correlates with Dengue, and Ear Bleeding is mapped to Hypertension.

### F. Clinical UI Overhaul (Visual Polish)
- Removed all development/port labels (e.g. Port 8001, Local Cache, Sprint 6) from the user interface.
- Added a pulsing EKG heartbeat waveform monitor (SVG path animation) to the right-hand panel when the system is idle.
- Added a radial gauge chart (animated progress ring) to display the primary diagnosis confidence percentage.
- Integrated specialist routing logic: Automatically recommends referring the patient to a Dermatologist, Cardiologist, Neurologist, etc., based on predicted outcomes.
- Implemented @media print CSS styles and a print button to generate clean Clinical Referral Reports (PDF) directly from the browser.

## 3. Architecture Overview
- React Client (Vercel) -> Asynchronous JSON Payloads -> FastAPI Backend (Render)
- FastAPI Backend -> TensorFlow-CPU (Model weights in RAM) -> Predictions & Probabilities
- FastAPI Backend -> motor (Asynchronous Driver) -> MongoDB Atlas Cloud Cluster
- React Client -> CSS Media Query Print Styles -> Referrals Diagnostic PDF Report

## 4. Cloud Deployments Summary
- Live Frontend URL: https://intelligent-diagnosis-system.vercel.app
- Live Backend API URL: https://mediwise-api.onrender.com
- Central Database: MongoDB Atlas Cluster0
"""

for line in content.split('\n'):
    if not line.strip():
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    else:
        doc.add_paragraph(line)

output_path = os.path.join('Sprint_7', 'Sprint7_Report.docx')
os.makedirs('Sprint_7', exist_ok=True)
doc.save(output_path)
print(f"Word document saved successfully to {output_path}!")
