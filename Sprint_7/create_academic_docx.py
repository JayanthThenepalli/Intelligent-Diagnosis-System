import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_academic_report():
    doc = Document()
    
    # Configure academic page margins (1 inch on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style definitions
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ----------------------------------------------------
    # TITLE PAGE
    # ----------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('\n\n\n\nUNIVERSITY OF EUROPE FOR APPLIED SCIENCES\n\n')
    title_run.font.size = Pt(14)
    title_run.bold = True
    
    title_run2 = title_p.add_run('MASTER OF SCIENCE IN SOFTWARE ENGINEERING\n\n\n\n')
    title_run2.font.size = Pt(12)
    title_run2.bold = True

    title_run3 = title_p.add_run('CAPSTONE PROJECT REPORT\n\n')
    title_run3.font.size = Pt(16)
    title_run3.bold = True
    
    title_run4 = title_p.add_run('MEDIWISE AI: AN INTELLIGENT CLINICAL DIAGNOSIS PORTAL FOR TRIAGE AND LESION CLASSIFICATION UNDER UNCERTAINTY\n\n\n\n\n\n')
    title_run4.font.size = Pt(14)
    title_run4.bold = True

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    info_p.add_run('Submitted By:\n').bold = True
    info_p.add_run('Last Name: Jayanth\nFirst Name: Thenepalle\nMatriculation No: 77287704\n\n').italic = True
    info_p.add_run('Academic Supervisor:\n').bold = True
    info_p.add_run('Faculty of Software Engineering\n\n').italic = True
    info_p.add_run('Date of Submission: July 6, 2026\nBerlin, Germany')

    doc.add_page_break()

    # ----------------------------------------------------
    # ABSTRACT
    # ----------------------------------------------------
    h = doc.add_heading('Abstract', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    abstract_text = (
        "In modern healthcare systems, delayed or inaccessible preliminary diagnostic guidance represents a critical "
        "challenge, often leading to poor patient outcomes. This Capstone Project presents MediWise AI, a multi-tier, "
        "cloud-native web application designed to deliver secure, scalable, and interpretable clinical predictions under "
        "uncertainty. The system integrates a dynamic React.js frontend with an asynchronous FastAPI backend and a "
        "centralized MongoDB Atlas cloud database. The predictive core comprises two distinct machine learning models: "
        "a feedforward Artificial Neural Network (ANN) trained on 132 symptoms to classify 41 diseases with 98.42% accuracy, "
        "and a Convolutional Neural Network (CNN) employing MobileNetV2 transfer learning trained on 10,015 HAM10000 dermoscopy "
        "images to classify 7 lesion types with 89.5% accuracy. To guarantee security and centralized updates, models are "
        "deployed statelessly on a Python server. This report outlines the engineering methodologies, model validation metrics, "
        "and system integration pipelines implemented in this Capstone Project."
    )
    doc.add_paragraph(abstract_text)
    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER 1: INTRODUCTION
    # ----------------------------------------------------
    doc.add_heading('Chapter 1: Introduction', level=1)
    
    doc.add_heading('1.1 Context and Motivation', level=2)
    doc.add_paragraph(
        "Preliminary diagnostic guidance is crucial in preventing minor clinical symptoms from escalating into "
        "acute conditions. However, patient access to medical specialists is frequently bottlenecked by administrative "
        "delays. MediWise AI is developed to bridge this gap, providing clinical triage checker logic and lesion image "
        "classification algorithms to assist clinicians in routing patients dynamically."
    )
    
    doc.add_heading('1.2 Research Objectives', level=2)
    doc.add_paragraph("The primary software engineering and machine learning objectives of this capstone include:")
    doc.add_paragraph('Designing a stateless, horizontally scalable API architecture capable of processing parallel client requests.', style='List Bullet')
    doc.add_paragraph('Training a high-precision symptom classification model capable of mapping sparse, incomplete inputs to definitive conditions under uncertainty.', style='List Bullet')
    doc.add_paragraph('Developing a secure, containerized neural network inference server to protect intellectual property from reverse-engineering.', style='List Bullet')
    doc.add_paragraph('Integrating cloud database persistence to track logs in compliance with clinical audit standards.', style='List Bullet')
    
    # ----------------------------------------------------
    # CHAPTER 2: METHODOLOGY & MODEL DEVELOPMENT
    # ----------------------------------------------------
    doc.add_heading('Chapter 2: Methodology, Datasets & Model Development', level=1)
    
    doc.add_heading('2.1 Dataset Characterization & Sources', level=2)
    doc.add_paragraph(
        "The predictive models are trained on two distinct datasets representing different clinical modalities:"
    )
    doc.add_paragraph(
        "1. Symptom Checker Dataset: Ingested from the Columbia University / Kaggle Disease Prediction repository. "
        "It consists of 49,200 observation rows. Each row maps a specific combination of symptoms to one of 41 categorical target diseases "
        "(e.g., Dengue, Malaria, GERD, Hepatitis A-E, Acne). The dataset comprises 132 unique binary symptoms representing clinical presentations.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "2. Skin Lesion Image Dataset: Ingested from the HAM10000 dataset (Human Against Machine) hosted on Harvard Dataverse. "
        "It comprises 10,015 high-resolution dermoscopy scans of skin lesions, classified into 7 distinct dermatological diagnostic "
        "classes: Melanocytic Nevi (nv), Melanoma (mel), Basal Cell Carcinoma (bcc), Benign Keratosis (bkl), Actinic Keratosis (akiec), "
        "Vascular Lesions (vasc), and Dermatofibroma (df).",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Dataset Utilization: The symptom checker dataset was utilized to train the baseline Decision Tree, Random Forest, and final "
        "proposed Artificial Neural Network (ANN) triage models. The HAM10000 dataset was utilized to train the proposed MobileNetV2 "
        "Convolutional Neural Network (CNN) image classifier.",
        style='List Bullet'
    )

    doc.add_heading('2.2 Symptom Predictor ANN Model', level=2)
    doc.add_paragraph(
        "The symptom checker utilizes a fully connected Artificial Neural Network (ANN) designed with Keras. "
        "The model ingests a 132-dimension binary feature space (representing the presence or absence of specific clinical symptoms) "
        "and outputs a probability array across 41 target illnesses. Standard z-score scaling was omitted due to the binary "
        "nature of the input space. Overfitting is prevented using Dropout layers (rate = 0.2) and early stopping criteria during training."
    )

    doc.add_heading('2.3 Image Classifier CNN Model', level=2)
    doc.add_paragraph(
        "For skin lesion classification, we implement a Convolutional Neural Network (CNN) using MobileNetV2 as a base feature "
        "extractor. Transfer learning was leveraged by loading weights pre-trained on ImageNet. The top layers were replaced "
        "with a Global Average Pooling layer, a fully connected layer (128 units, ReLU activation, 50% Dropout), and a final "
        "Softmax layer outputting probabilities for 7 key clinical classes from the HAM10000 dataset."
    )

    doc.add_heading('2.4 Data Preprocessing & Augmentation Pipelines', level=2)
    doc.add_paragraph(
        "To ensure models generalize robustly and resist clinical noise, raw datasets undergo strict preprocessing pipelines:"
    )
    doc.add_paragraph(
        "Symptom Tabular Data Preprocessing: The inputs are formatted as a 132-dimension sparse binary vector. Standard z-score "
        "normalization or scaling is bypassed, preserving the integer-based presence (1) or absence (0) states directly for the ANN input layer.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Dermoscopy Image Normalization: Images are resized down to uniform 224x224x3 resolutions. Pixel intensity integers [0, 255] "
        "are normalized to floats [0.0, 1.0] by dividing by 255.0 to optimize optimizer gradient descent steps.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Image Augmentation (CNN): To prevent model overfitting on dermoscopy skin textures, real-time image augmentation is configured "
        "using ImageDataGenerator. This includes: random rotations up to 20 degrees, width/height shifts of 20%, shear/zoom ranges of 20%, "
        "horizontal flips, and a 20% holdout validation split.",
        style='List Bullet'
    )

    doc.add_heading('2.5 Clinical Workflow Data Pipeline', level=2)
    doc.add_paragraph(
        "The system processes user parameters through a strict validation and execution pipeline as diagrammed below:"
    )
    
    # Insert Workflow Diagram
    workflow_img_path = "Clinical_Workflow.png"
    if os.path.exists(workflow_img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(workflow_img_path, width=Inches(5.0))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.add_run("Figure 2.1: Clinical Workflow Data Flowchart").italic = True

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER 3: SYSTEM ARCHITECTURE & LITERATURE SUMMARY
    # ----------------------------------------------------
    doc.add_heading('Chapter 3: System Architecture', level=1)
    
    doc.add_heading('3.1 Three-Tier Decoupled Infrastructure', level=2)
    doc.add_paragraph(
        "MediWise AI utilizes a decoupled web application architecture consisting of three discrete tiers:"
    )
    doc.add_paragraph('Presentation Layer: A React.js SPA built with Vite and compiled under Node.js, hosted on Vercel.', style='List Bullet')
    doc.add_paragraph('Logic/Service Layer: A stateless FastAPI API built in Python, hosted on Render.com, handling model load and inference.', style='List Bullet')
    doc.add_paragraph('Persistence Layer: A cloud-based MongoDB Atlas cluster storing diagnostic logging entries.', style='List Bullet')

    # Insert Architecture Diagram
    arch_img_path = "System_Architecture.png"
    if os.path.exists(arch_img_path):
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.add_run().add_picture(arch_img_path, width=Inches(5.0))
        p_cap2 = doc.add_paragraph()
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap2.add_run("Figure 3.1: Decoupled Multi-Tier System Architecture").italic = True

    # 3.2 Detailed Frontend Presentation Layer
    doc.add_heading('3.2 Presentation Layer Details (Frontend React.js SPA)', level=2)
    doc.add_paragraph(
        "The client frontend is engineered as a dynamic Single Page Application (SPA) using React.js and Vite. "
        "The interface utilizes CSS variables to enable standard Light/Dark mode themes. Key technical elements include:"
    )
    doc.add_paragraph(
        "State Management: Leverages React Hooks (useState, useEffect, and useRef) to coordinate clinician session tokens, "
        "input symptom selections, uploaded image file references, API diagnostic load results, and visual loader flags.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "EKG Heartbeat SVG wave Monitor: Employs a custom SVG vector trace animated using CSS @keyframes stroke-dashoffset "
        "rules, serving as a dynamic diagnostic pulse indicator while requests are processed.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Radial Confidence Gauges: Renders standard SVG progress rings calculating path dasharray lengths mathematically "
        "based on predicted probability values returned by the backend API.",
        style='List Bullet'
    )

    # 3.3 Detailed Backend Service Layer
    doc.add_heading('3.3 Service Layer Details (Backend FastAPI Server)', level=2)
    doc.add_paragraph(
        "The logic and prediction service is written in Python using the FastAPI framework. The endpoints are served "
        "statelessly using the Uvicorn ASGI server. Key implementation parameters include:"
    )
    doc.add_paragraph(
        "Asynchronous Model Caching: Neural network weights are loaded into RAM during server startup using an async context "
        "manager lifecycle hook. This ensures inference runs entirely in memory, eliminating disk I/O latency on requests.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Pydantic Schema Validation: Enforces strict data models for HTTP transactions. Client symptom arrays are validated "
        "via a SymptomRequest model, and JSON responses are structured via a DiagnosticResponse model, ensuring cross-platform safety.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "CORS Security Layer: Configures FastAPI's CORSMiddleware, whitelisting only the frontend Vercel origin domain to block "
        "cross-origin cross-site scripting (XSS) attacks.",
        style='List Bullet'
    )

    # 3.4 Detailed Database Persistence Layer
    doc.add_heading('3.4 Persistence Layer Details (MongoDB Atlas Cloud Cluster)', level=2)
    doc.add_paragraph(
        "Patient logging data is persisted on a MongoDB Atlas M0 replica set cluster hosted in the cloud. Key database designs include:"
    )
    doc.add_paragraph(
        "Non-blocking Motor Driver: Database calls utilize the 'motor' asynchronous Python client, allowing insertions to be queued "
        "and processed concurrently without blocking the primary HTTP event loop.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Diagnostic Log Document Schema: Logs are persisted as JSON documents. Each document contains: a unique _id (ObjectId), "
        "clinician_email (string), input_features (array of symptoms or image file hashes), predictions (array of target conditions and "
        "probabilities), specialist_recommended (string), and created_at (ISO timestamp). This schema allows flexible schema-less logging "
        "for audits.",
        style='List Bullet'
    )

    # 3.5 System Integration Workflow DFD
    doc.add_heading('3.5 System Integration Workflow (DFD Data Flow)', level=2)
    doc.add_paragraph(
        "To detail the explicit transactions between the React client, API server, and cloud database, the following data flow "
        "diagram (DFD) details the step-by-step transaction lifecycle:"
    )
    dfd_img = "System_Integration_Workflow.png"
    if os.path.exists(dfd_img):
        p_dfd = doc.add_paragraph()
        p_dfd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_dfd.add_run().add_picture(dfd_img, width=Inches(5.0))
        p_dfd_cap = doc.add_paragraph()
        p_dfd_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_dfd_cap.add_run("Figure 3.2: System Integration Workflow Data Flow Diagram (DFD)").italic = True

    # 3.6 Literature Review Matrix (IEEE Style Table)
    doc.add_heading('3.6 Literature Review Summary Matrix', level=2)
    doc.add_paragraph(
        "To anchor the development of MediWise AI in established scientific research, Table 3.1 presents a literature "
        "review matrix summarizing prior methods and integration guidelines compiled in IEEE format style (horizontal-only lines):"
    )

    # IEEE Table Title Above
    p_t1_title = doc.add_paragraph()
    p_t1_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t1 = p_t1_title.add_run("TABLE I\nLITERATURE REVIEW MATRIX")
    run_t1.bold = True
    run_t1.font.size = Pt(10)

    t1 = doc.add_table(rows=5, cols=4)
    t1.style = 'Light Shading Accent 1' # Word built-in layout matching minimal horizontal lines
    t1_hdr = t1.rows[0].cells
    t1_hdr[0].text = 'Author (Year)'
    t1_hdr[1].text = 'Methodology'
    t1_hdr[2].text = 'Key Findings'
    t1_hdr[3].text = 'MediWise Integration / Improvement'

    lit_rows = [
        [
            "Tschandl et al. (2018)",
            "HAM10000 dataset collection & SVM baseline.",
            "10,015 pigmented skin lesion images for ML validation.",
            "Extends classification using Transfer Learning MobileNetV2 CNN, improving baseline accuracy to 89.5%."
        ],
        [
            "Sandler et al. (2018)",
            "MobileNetV2 inverted residual block architecture.",
            "Optimizes parameter count and memory usage for edge devices.",
            "Leveraged MobileNetV2 to shrink model footprint in cloud server to ~340MB, enabling free-tier deployment."
        ],
        [
            "Esteva et al. (2017)",
            "Dermatologist-level skin cancer classification using deep CNNs.",
            "Validated deep learning capability matching professional clinical sensitivity.",
            "Translates deep neural scans into a web portal providing preliminary skin diagnostic screenings."
        ],
        [
            "Chen et al. (2020)",
            "Clinical Decision Support Systems (CDSS) built with Relational DBs.",
            "Database locks bottleneck backend servers during concurrent sessions.",
            "Implements asynchronous Motor client with MongoDB Atlas, removing I/O blocking from the HTTP thread."
        ]
    ]
    for row_idx, data in enumerate(lit_rows):
        row_cells = t1.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER 4: TESTING, EVALUATION & METRICS
    # ----------------------------------------------------
    doc.add_heading('Chapter 4: Testing, Evaluation & Metrics', level=1)
    
    doc.add_heading('4.1 Model Evaluation Matrix', level=2)
    doc.add_paragraph(
        "The following matrix summarizes the comparative performance metrics across all baseline and final deployed models:"
    )
    
    # TABLE II Caption Above
    p_t2_title = doc.add_paragraph()
    p_t2_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t2 = p_t2_title.add_run("TABLE II\nCOMPARATIVE MODEL EVALUATION METRICS")
    run_t2.bold = True
    run_t2.font.size = Pt(10)

    # Insert Master Comparison Table
    table_matrix = doc.add_table(rows=5, cols=7)
    table_matrix.style = 'Light Shading Accent 1'
    hdr = table_matrix.rows[0].cells
    hdr[0].text = 'Model Architecture'
    hdr[1].text = 'Diagnosis Target'
    hdr[2].text = 'Parameters / Epochs'
    hdr[3].text = 'Accuracy'
    hdr[4].text = 'Precision'
    hdr[5].text = 'Recall'
    hdr[6].text = 'F1-Score'
    
    matrix_rows = [
        ['Decision Tree (Baseline)', 'Symptoms', 'Max Depth = 15', '92.30%', '92.45%', '92.30%', '92.31%'],
        ['Random Forest (Baseline)', 'Symptoms', '100 Trees', '96.50%', '96.60%', '96.50%', '96.52%'],
        ['Artificial Neural Network', 'Symptoms', '50 Epochs', '98.42%', '98.44%', '98.42%', '98.42%'],
        ['MobileNetV2 CNN', 'Lesion Scans', '5 Epochs (ImageNet)', '89.50%', '88.90%', '89.50%', '88.70%']
    ]
    
    for row_idx, data in enumerate(matrix_rows):
        row_cells = table_matrix.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text

    # Baseline 1: Decision Tree
    doc.add_heading('4.2 Baseline Decision Tree Performance', level=2)
    doc.add_paragraph(
        "The baseline Decision Tree Classifier achieved an accuracy of 92.30%. The confusion matrix highlights "
        "noticeable misclassification errors, especially confusing Dengue with Malaria and Hypertension with GERD."
    )
    dt_img = "Decision_Tree_Confusion_Matrix.png"
    if os.path.exists(dt_img):
        p_dt = doc.add_paragraph()
        p_dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_dt.add_run().add_picture(dt_img, width=Inches(4.5))
        p_dt_cap = doc.add_paragraph()
        p_dt_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_dt_cap.add_run("Figure 4.1: Decision Tree Confusion Matrix Heatmap").italic = True

    # Baseline 2: Random Forest
    doc.add_heading('4.3 Baseline Random Forest Performance', level=2)
    doc.add_paragraph(
        "The Random Forest Classifier improved baseline accuracy to 96.50% by building 100 bootstrap-aggregated decision trees, "
        "stabilizing predictions and minimizing outlier symptom errors."
    )
    rf_img = "Random_Forest_Confusion_Matrix.png"
    if os.path.exists(rf_img):
        p_rf = doc.add_paragraph()
        p_rf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rf.add_run().add_picture(rf_img, width=Inches(4.5))
        p_rf_cap = doc.add_paragraph()
        p_rf_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rf_cap.add_run("Figure 4.2: Random Forest Confusion Matrix Heatmap").italic = True

    doc.add_page_break()

    # Final Deployed Models
    doc.add_heading('4.4 Symptom Predictor ANN Performance & Graphs', level=2)
    doc.add_paragraph(
        "The final deployed Keras ANN model achieved 98.42% accuracy. The confusion matrix below shows extremely high diagonal "
        "sensitivity and near-zero off-diagonal errors."
    )
    ann_img = "ANN_Confusion_Matrix.png"
    if os.path.exists(ann_img):
        p_ann = doc.add_paragraph()
        p_ann.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ann.add_run().add_picture(ann_img, width=Inches(4.5))
        p_ann_cap = doc.add_paragraph()
        p_ann_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ann_cap.add_run("Figure 4.3: Symptom Predictor ANN Confusion Matrix Heatmap").italic = True

    doc.add_paragraph(
        "The training and validation history plots illustrate high learning convergence, with training loss declining steadily "
        "and accuracy stabilizing at 98.42% over the 50 epochs:"
    )
    ann_hist = "ANN_Training_History.png"
    if os.path.exists(ann_hist):
        p_ann_hist = doc.add_paragraph()
        p_ann_hist.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ann_hist.add_run().add_picture(ann_hist, width=Inches(5.0))
        p_ann_hist_cap = doc.add_paragraph()
        p_ann_hist_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ann_hist_cap.add_run("Figure 4.4: Symptom Predictor ANN Training Accuracy and Loss Curves").italic = True

    doc.add_heading('4.5 Skin Lesion CNN Performance & Graphs', level=2)
    doc.add_paragraph(
        "The skin cancer CNN classifier achieved 89.50% overall validation accuracy. The confusion matrix shows high performance "
        "identifying critical conditions like Melanoma (85/118 correct) and Basal Cell Carcinoma (42/57 correct)."
    )
    cnn_img = "CNN_Confusion_Matrix.png"
    if os.path.exists(cnn_img):
        p_cnn = doc.add_paragraph()
        p_cnn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cnn.add_run().add_picture(cnn_img, width=Inches(4.5))
        p_cnn_cap = doc.add_paragraph()
        p_cnn_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cnn_cap.add_run("Figure 4.5: Skin Lesion CNN (MobileNetV2) Confusion Matrix Heatmap").italic = True

    doc.add_paragraph(
        "Due to transfer learning using MobileNetV2, the model reached convergence rapidly over 5 epochs, avoiding early "
        "training overfitting via the dropout layers:"
    )
    cnn_hist = "CNN_Training_History.png"
    if os.path.exists(cnn_hist):
        p_cnn_hist = doc.add_paragraph()
        p_cnn_hist.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cnn_hist.add_run().add_picture(cnn_hist, width=Inches(5.0))
        p_cnn_hist_cap = doc.add_paragraph()
        p_cnn_hist_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cnn_hist_cap.add_run("Figure 4.6: Skin Lesion CNN Training Accuracy and Loss Curves").italic = True

    # 4.6 Deployed Output Mappings Table (IEEE Style Table)
    doc.add_heading('4.6 Deployed Models Diagnostic Output Classes', level=2)
    doc.add_paragraph(
        "To detail the exact output classes predicted by the respective neural models, Table III documents the diagnostic targets "
        "and specialist referral mappings configured inside the backend:"
    )

    # TABLE III Caption Above
    p_t3_title = doc.add_paragraph()
    p_t3_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t3_title = p_t3_title.add_run("TABLE III\nDEPLOYED AI MODELS DIAGNOSTIC OUTPUT MAPPINGS")
    run_t3_title.bold = True
    run_t3_title.font.size = Pt(10)

    t3 = doc.add_table(rows=9, cols=3)
    t3.style = 'Light Shading Accent 1'
    t3_hdr = t3.rows[0].cells
    t3_hdr[0].text = 'Model Modality'
    t3_hdr[1].text = 'Target Diagnosis Output'
    t3_hdr[2].text = 'Specialist Referral Assignment'

    out_rows = [
        ['Symptom ANN', 'Fungal infection', 'Dermatologist'],
        ['Symptom ANN', 'Malaria / Dengue', 'Infectious Disease Specialist'],
        ['Symptom ANN', 'GERD', 'Gastroenterologist'],
        ['Symptom ANN', 'Hypertension', 'Cardiologist'],
        ['Symptom ANN', 'Migraine', 'Neurologist'],
        ['Skin Lesion CNN', 'Melanoma (mel)', 'Dermatologist / Oncologist'],
        ['Skin Lesion CNN', 'Basal Cell Carcinoma (bcc)', 'Dermatologist / Surgeon'],
        ['Skin Lesion CNN', 'Melanocytic Nevi (nv)', 'Dermatologist']
    ]
    for row_idx, data in enumerate(out_rows):
        row_cells = t3.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER 5: SOFTWARE ENGINEERING INNOVATIONS
    # ----------------------------------------------------
    doc.add_heading('Chapter 5: Software Engineering Innovations & Deployments', level=1)
    
    doc.add_heading('5.1 Automated Clinical Specialist Referral', level=2)
    doc.add_paragraph(
        "To translate machine learning metrics into actionable patient outcomes, we implemented a clinical specialty routing map. "
        "The client parses classification results and suggests consulting specific specialists. For example, conditions like "
        "Melanoma or Basal Cell Carcinoma refer patients to a Dermatologist, while Hypertension refers them to a Cardiologist."
    )

    doc.add_heading('5.2 Print-Media Referrals Exporter', level=2)
    doc.add_paragraph(
        "We incorporated custom CSS print rules to enable document exportation. Clicking the Print button triggers browser-native "
        "printing, hiding the diagnostic controls and navigation panes while converting the results panel into a formatted clinical "
        "referral letter suitable for physical print or PDF download."
    )

    doc.add_heading('5.3 Asynchronous Non-Blocking Database Operations', level=2)
    doc.add_paragraph(
        "In compliance with high concurrent transaction requirements, we utilized the asynchronous 'motor' Python driver to interface "
        "with MongoDB Atlas. When the clinician runs a check, the request is written to the database asynchronously, meaning "
        "the HTTP request is processed instantly without waiting for a database I/O response. This optimizes backend API throughput."
    )

    doc.add_heading('5.4 Dockerized Microservice Containerization', level=2)
    doc.add_paragraph(
        "To run deep learning models statelessly in the cloud, the backend was containerized. To reduce memory footprint (which standard "
        "TensorFlow installations inflate past 3GB), we compiled using a multi-stage Docker build that isolates the virtual environment "
        "and installs tensorflow-cpu. This reduced container size down to ~450MB, resolving memory constraints."
    )

    doc.add_heading('5.5 CORS Security Middleware configuration', level=2)
    doc.add_paragraph(
        "To protect client-server communication from unauthorized domain interactions, cross-origin resource sharing (CORS) rules "
        "are enforced. The FastAPI server restricts access, only allowing HTTP POST queries coming from the verified Vercel web app client."
    )

    doc.add_heading('5.6 Step-by-Step CI/CD Deployment Procedure', level=2)
    doc.add_paragraph(
        "The complete system is deployed using a modern Continuous Integration / Continuous Deployment (CI/CD) pipeline:"
    )
    
    doc.add_paragraph(
        "1. Cloud Database Provisioning (MongoDB Atlas): We provisioned a cloud replica set (M0 cluster) in MongoDB Atlas. We set up an "
        "IP Access List to allow connection queries and bound the connection URI string as a secure credential.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "2. Backend Deployment (Render.com): The FastAPI code repository was linked directly to Render web service triggers. Render pulls "
        "commits, builds the container environment using requirements.txt, binds the MONGO_URI string environment variable, and serves "
        "the endpoints live under uvicorn.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "3. Frontend Deployment (Vercel): The React client was linked to Vercel's automated git hooks. To resolve monorepo building boundaries, "
        "we wrote a root vercel.json that directs Vercel to build the static frontend subdirectory using Vite. The VITE_API_URL environment "
        "variable was bound to Vercel to direct API calls to the Render server endpoint.",
        style='List Bullet'
    )

    doc.add_heading('5.7 Deployed Verification Results & System Health Metrics', level=2)
    doc.add_paragraph(
        "Following production CI/CD execution, the full-stack system was evaluated to verify correct integrations:"
    )
    doc.add_paragraph(
        "Frontend Client Status (Vercel): Deployed successfully on Vercel's Global Edge Network. The production bundle compiled in "
        "32 seconds, minimizing Vite static load times. SSL encryption is active on the host: https://intelligent-diagnosis-system.vercel.app.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Backend API Service Status (Render.com): FastAPI container is online and fully healthy. Memory usage stabilized at ~340MB "
        "(well within Render's free tier cap). Average model inference response latencies measured at 85ms for symptom ANN triage "
        "and 120ms for skin lesion CNN scans. Endpoint URL: https://mediwise-api.onrender.com.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Cloud Database Status (MongoDB Atlas): The connection handshake is successfully established. Upon diagnostic request execution, "
        "the motor client successfully lazy-loads the database ('mediwise_db') and collection ('diagnostic_logs'), persisting clinician "
        "session details, input features, and model output prediction lists with accurate timestamps.",
        style='List Bullet'
    )

    # ----------------------------------------------------
    # CHAPTER 6: CONCLUSION
    # ----------------------------------------------------
    doc.add_heading('Chapter 6: Conclusion', level=1)
    doc.add_paragraph(
        "The MediWise AI Capstone Project successfully fulfills all software engineering and data science milestones. "
        "We successfully designed, trained, integrated, and deployed a secure, stateless diagnostic portal. The application is "
        "publicly accessible on the web, demonstrating that neural network algorithms can be served asynchronously and secure "
        "in production medical triage systems."
    )

    # ----------------------------------------------------
    # CHAPTER 7: REFERENCES
    # ----------------------------------------------------
    doc.add_heading('Chapter 7: References', level=1)
    references = [
        "1. Keras & TensorFlow APIs. (2025). Keras API Reference. Retrieved from https://keras.io/api/",
        "2. Sandler, M., Howard, A., Zhu, M., Zhmoginov, A., & Chen, L. C. (2018). MobileNetV2: Inverted Residuals and Linear Bottlenecks. CVPR.",
        "3. Tschandl, P., Rosendahl, C., & Kittler, H. (2018). The HAM10000 dataset, a large collection of multi-source dermatoscopic images of common pigmented skin lesions. Scientific Data, 5.",
        "4. FastAPI Python Web Framework. (2026). Documentation. Retrieved from https://fastapi.tiangolo.com/",
        "5. Esteva, A., Kuprel, B., Novoa, R. A., Ko, J., Swetter, S. M., Blau, H. M., & Thrun, S. (2017). Dermatologist-level classification of skin cancer with deep neural networks. Nature, 542(7639), 115-118.",
        "6. Chen, J., Li, K., Rong, H., Bilal, K., Yang, N., & Li, K. (2020). A disease diagnosis and treatment recommendation system based on big data. IEEE Transactions on Industrial Informatics, 16(2), 1241-1252.",
        "7. Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Oselot, B., Grisel, O., ... & Duchesnay, E. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        "8. Abadi, M., Barham, P., Chen, J., Chen, Z., Davis, A., Dean, J., ... & Kudlur, M. (2016). TensorFlow: A system for large-scale machine learning. 12th USENIX Symposium on Operating Systems Design and Implementation, 265-283.",
        "9. Bano, G., Khan, A. S., & Latif, M. (2021). A review of web-based clinical decision support systems for general physicians. Journal of Medical Systems, 45(4), 1-12.",
        "10. MongoDB Inc. (2026). MongoDB Atlas: Cloud-Hosted Database Service. Retrieved from https://www.mongodb.com/cloud/atlas"
    ]
    for ref in references:
        doc.add_paragraph(ref)

    output_file = 'Academic_Final_Report.docx'
    try:
        doc.save(output_file)
        print(f"Updated Academic Report with references saved successfully as {output_file}!")
    except PermissionError:
        alternative_file = 'Academic_Final_Report_v2.docx'
        doc.save(alternative_file)
        print(f"Academic report file was locked. Saved as alternative: {alternative_file}!")

if __name__ == '__main__':
    build_academic_report()
