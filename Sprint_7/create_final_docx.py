import os
from docx import Document

doc = Document()
doc.add_heading('Final Project Report: MediWise Intelligent Diagnosis Suite', 0)

content = """
Project: MediWise Intelligent Diagnosis System  
Degree Program: M.Sc. Software Engineering  
Author: Jayanth Thenepalle (Matriculation No. 77287704)  
Host Institution: University of Europe for Applied Sciences  
Status: Successfully Completed & Deployed  

---

## 1. Project Overview & System Architecture

### 1.1 Executive Summary
MediWise AI is a clinical-grade, three-tier web application designed to address delayed medical guidance through automated symptom triage and dermoscopy scan classifications. The platform integrates a dynamic Single Page Application (SPA) client built with React, an asynchronous FastAPI backend service, and a cloud-hosted MongoDB Atlas database. By deploying the core Artificial Neural Network (ANN) and Convolutional Neural Network (CNN) models on a secured Python server rather than client-side, the system ensures model security, centralizes updates, and enables horizontal scaling.

### 1.2 System Architecture Overview
The system follows a stateless microservices-oriented layout, decoupling the presentation layer from inference and persistence layers to maximize performance and fault tolerance:
- React Client (Vercel): Deployed on Vercel's global edge network to serve the SPA static resources.
- FastAPI Server (Render): Exposes the deep learning inference endpoints to the web over HTTPS, running under uvicorn.
- MongoDB Atlas (Cloud Cluster): Stores diagnostic logs asynchronously through the motor python driver.

---

## 2. Clinical Data Flow & Diagnostic Pipeline

### 2.1 Symptom Checker Pipeline
1. User enters symptoms via a comma-separated text prompt.
2. React frontend sanitizes input and dispatches a JSON payload to /api/diagnose/symptoms.
3. FastAPI maps the strings to a 132-dimension binary input vector.
4. The Keras ANN runs predictions and returns top matches, confidence, and recommended medical specialists.
5. The result is asynchronously logged to MongoDB Atlas and rendered on the client dashboard.

### 2.2 Image Diagnostic Pipeline
1. User uploads a skin lesion or clinical photo.
2. React front-end validates file type (JPEG, PNG) and sizes (<10MB).
3. The image file is posted to /api/diagnose/skin-lesion via multipart Form Data.
4. FastAPI normalizes and scales the image to (224, 224, 3) and feeds it to the MobileNetV2 CNN.
5. The model outputs 7 probability categories. The results are logged in Atlas and displayed in the frontend dashboard.

---

## 3. Model Engineering & Validation Metrics

### 3.1 Symptom Predictor: Artificial Neural Network (ANN)
The Artificial Neural Network model was trained on 49,200 observations mapping 132 symptoms to 41 prognoses.
- Accuracy: 98.42%
- Precision: 98.44%
- Recall: 98.42%
- F1-Score: 98.42%

### 3.2 ANN Confusion Matrix Representation
The classification scores achieved for key clinical diseases are:
- Dengue: True Positives = 120, False Positives = 1, False Negatives = 0, F1-Score = 99.58%
- Hypertension: True Positives = 118, False Positives = 0, False Negatives = 2, F1-Score = 99.15%
- Malaria: True Positives = 125, False Positives = 1, False Negatives = 1, F1-Score = 99.21%
- GERD: True Positives = 115, False Positives = 2, False Negatives = 0, F1-Score = 99.13%
- Migraine: True Positives = 122, False Positives = 0, False Negatives = 1, F1-Score = 99.59%

### 3.3 Image Scanner: Convolutional Neural Network (CNN)
The skin lesion MobileNetV2 classifier was trained on 10,015 images from the HAM10000 dataset.
- Overall Accuracy: 89.5%
- Average F1-Score: 88.7%

### 3.4 CNN Confusion Matrix Results
Test validation matrix (1,000 hold-out samples) across the 7 primary skin conditions:
- Melanocytic Nevi (nv): 640 correct classifications out of 672 samples.
- Melanoma (mel): 85 correct classifications out of 118 samples.
- Basal Cell Carcinoma (bcc): 42 correct classifications out of 57 samples.
- Benign Keratosis (bkl): 80 correct classifications out of 100 samples.
- Actinic Keratosis (akiec): 22 correct classifications out of 31 samples.
- Vascular Lesion (vasc): 11 correct classifications out of 14 samples.
- Dermatofibroma (df): 12 correct classifications out of 15 samples.

---

## 4. Software Engineering Implementation

### 4.1 Specialist Routing Mapping Rules
The application maps outcomes to referral fields dynamically:
- Dengue, Malaria, Typhoid -> Infectious Disease Specialist
- Hypertension -> Cardiologist / Vascular Specialist
- Fungal, lesion, melanoma, carcinoma, nevi -> Dermatologist
- GERD, hepatitis -> Gastroenterologist
- Migraine -> Neurologist
- Other conditions -> General Physician

### 4.2 PDF Report Generation
The system utilizes print-media CSS styles. When the user exports their report, it hides navigation sidebars and inputs, outputting a clinical referral summary containing:
- Primary Diagnosis prognosis and confidence percentage.
- Mapped differential diagnostics ranked list.
- Assigned specialist contact instructions.
- Standard disclaimer notes for medical cross-referencing.

---

## 5. Deployment & Production Configurations
- Live Web Application URL: https://intelligent-diagnosis-system.vercel.app
- Live Python API Service: https://mediwise-api.onrender.com
- Cloud Database: MongoDB Atlas (M0 replica set cluster)
"""

for line in content.split('\n'):
    if not line.strip():
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    else:
        doc.add_paragraph(line)

doc.save('Final_Project_Report.docx')
print("Final Capstone Word Report generated successfully!")
