import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_main_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.bold = True
    run.underline = True
    return p

def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.bold = True
    return p

def add_body_text(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(11)  # Blank line space after paragraph
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def add_bullet_point(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def add_figure(doc, img_path, caption):
    if os.path.exists(img_path):
        # One blank line space before
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(6)
        
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(6)
        p_img.add_run().add_picture(img_path, width=Inches(4.5))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        run = p_cap.add_run(caption)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        
        # One blank line space after
        p_space2 = doc.add_paragraph()
        p_space2.paragraph_format.space_after = Pt(6)

def build_report():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        
    # Set default style to Calibri 11pt
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)

    # ----------------------------------------------------
    # PAGE 1: TITLE PAGE
    # ----------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Large spacing before name
    name_run = title_p.add_run('\n\n\n\n\n\nJayanth Thenepalle\n\n')
    name_run.font.size = Pt(14)
    name_run.bold = True
    
    title_run = title_p.add_run('MEDIWISE AI: AN INTELLIGENT CLINICAL DIAGNOSIS PORTAL FOR TRIAGE AND LESION CLASSIFICATION UNDER UNCERTAINTY\n\n')
    title_run.font.size = Pt(14)
    title_run.bold = True

    project_run = title_p.add_run('EBTM 881 CAPSTONE PROJECT\n\n')
    project_run.font.size = Pt(12)
    project_run.bold = True

    advisor_run = title_p.add_run('ADVISOR: Dr. Faculty Supervisor')
    advisor_run.font.size = Pt(12)
    advisor_run.bold = True
    
    doc.add_page_break()

    # ----------------------------------------------------
    # PAGE 2: TABLE OF CONTENTS
    # ----------------------------------------------------
    toc_p = doc.add_paragraph()
    toc_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_run = toc_p.add_run('Table of Contents\n\n')
    toc_run.font.size = Pt(14)
    toc_run.bold = True
    
    toc_items = [
        "Table of Contents.........................................................................................................................................2",
        "1. Introduction & Problem Motivation........................................................................................................3",
        "   1.1 Section Headings & Bullets............................................................................................................3",
        "   1.2 Figures and Tables..........................................................................................................................4",
        "2. Problem Statement..................................................................................................................................4",
        "3. Background & Literature Review.............................................................................................................5",
        "4. Data...........................................................................................................................................................5",
        "5. Model and Analysis..................................................................................................................................5",
        "6. Results and Recommendations...............................................................................................................5",
        "7. Conclusions...............................................................................................................................................5",
        "8. Acknowledgements..................................................................................................................................5",
        "9. References................................................................................................................................................5"
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        
    doc.add_page_break()

    # ----------------------------------------------------
    # 1. INTRODUCTION & PROBLEM MOTIVATION
    # ----------------------------------------------------
    add_main_heading(doc, "1. Introduction & Problem Motivation")
    add_body_text(
        doc,
        "In modern healthcare settings, delayed or inaccessible preliminary diagnostic guidance represents a critical "
        "challenge, often leading to poor patient outcomes. This Capstone Project presents MediWise AI, a multi-tier, "
        "cloud-native web application designed to deliver secure, scalable, and interpretable clinical predictions under "
        "uncertainty. The system integrates a dynamic React.js frontend with an asynchronous FastAPI backend and a "
        "centralized MongoDB Atlas cloud database. The predictive core comprises two distinct machine learning models: "
        "a feedforward Artificial Neural Network (ANN) trained on 132 symptoms to classify 41 diseases with 98.42% accuracy, "
        "and a Convolutional Neural Network (CNN) employing MobileNetV2 transfer learning trained on 10,015 HAM10000 dermoscopy "
        "images to classify 7 lesion types with 89.5% accuracy."
    )
    add_body_text(
        doc,
        "The primary software engineering and machine learning objectives of this capstone include designing a stateless, "
        "horizontally scalable API architecture capable of processing parallel client requests while maintaining model security "
        "and cloud persistence. By deploying the core neural network models on a secured Python server rather than client-side, "
        "the system ensures model security, centralizes updates, and enables horizontal scaling."
    )

    add_sub_heading(doc, "1.1 Section Headings & Bullets")
    add_body_text(
        doc,
        "This subsection outlines the structural conventions applied to ensure consistency throughout this capstone document. "
        "Specific bullet guidelines are outlined below:"
    )
    add_bullet_point(doc, "First level bullet.", level=0)
    add_bullet_point(doc, "Second level bullet.", level=1)
    add_bullet_point(doc, "Third level bullet.", level=2)

    add_sub_heading(doc, "1.2 Figures and Tables")
    add_body_text(
        doc,
        "Figures and tables are integrated directly within the text flow to verify system layout implementations. "
        "Figure 1.1 illustrates the clinical data pipeline flowchart mapping symptom and image checker transactions. "
        "Figure 1.2 illustrates the decoupled three-tier system architecture."
    )
    
    add_figure(doc, "Clinical_Workflow.png", "Figure 1: Example clinical workflow data pipeline flowchart")
    add_figure(doc, "System_Architecture.png", "Figure 2: Example system architecture layout")

    # ----------------------------------------------------
    # 2. PROBLEM STATEMENT
    # ----------------------------------------------------
    add_main_heading(doc, "2. Problem Statement")
    add_body_text(
        doc,
        "Medical triage systems frequently struggle to handle patient inputs under clinical uncertainty constraints. "
        "These constraints manifest as sparse or incomplete symptom checklists, clinical noise, and subjective patient "
        "descriptions. In addition, client-side model execution exposes proprietary network weights to security vulnerabilities "
        "and is restricted by local client computing resources."
    )
    add_body_text(
        doc,
        "To address these issues, this project implements a highly secure, server-side neural network inference engine "
        "integrated with a decoupled React client and MongoDB logging persistence. The scope of analysis covers: tabular symptom "
        "triaging across 41 classes, automated dermoscopic scan classification across 7 skin lesion types, asynchronous database logging, "
        "and clinical care specialist referrals."
    )

    # ----------------------------------------------------
    # 3. BACKGROUND & LITERATURE REVIEW
    # ----------------------------------------------------
    add_main_heading(doc, "3. Background & Literature Review")
    add_body_text(
        doc,
        "Prior research in healthcare automation highlights the utility of deep learning models for classification tasks. "
        "For tabular symptom mapping, feedforward Artificial Neural Networks (ANNs) have demonstrated superior precision "
        "compared to traditional decision trees due to their capacity to capture non-linear feature interactions. "
        "For image classification, MobileNetV2 stands out as an efficient Convolutional Neural Network (CNN) architecture, "
        "employing depthwise separable convolutions to maintain high performance with a low parameter footprint. "
        "This project builds on these models, deploying them statelessly using FastAPI to support scale."
    )

    # ----------------------------------------------------
    # 4. DATA
    # ----------------------------------------------------
    add_main_heading(doc, "4. Data")
    add_body_text(
        doc,
        "The Capstone Project leverages two clinical datasets for training and validation:"
    )
    add_bullet_point(
        doc,
        "Symptom Dataset: Ingested from the Columbia University / Kaggle Disease Prediction dataset. It comprises "
        "49,200 observation rows mapping 132 sparse binary symptoms to 41 target diseases. Normalization is bypassed to preserve binary sparsity.",
        level=0
    )
    add_bullet_point(
        doc,
        "Skin Lesion Dataset: Ingested from the HAM10000 (Human Against Machine) dataset hosted on Harvard Dataverse, containing "
        "10,015 dermoscopy scans. Images are normalized to 224x224x3 resolutions and pixel intensities are rescaled from [0, 255] to [0.0, 1.0].",
        level=0
    )
    add_body_text(
        doc,
        "Image augmentations (rotations, shifts, flips, and shear/zoom) are applied in real-time to prevent model overfitting. "
        "A 20% holdout validation partition is utilized to monitor accuracy during training."
    )

    # ----------------------------------------------------
    # 5. MODEL AND ANALYSIS
    # ----------------------------------------------------
    add_main_heading(doc, "5. Model and Analysis")
    add_body_text(
        doc,
        "The clinical application is engineered using a three-tier decoupled structure. The frontend is a React.js SPA "
        "utilizing state hooks and animated SVGs to present a clinic-ready console. The FastAPI backend loads trained model "
        "weights into RAM during startup to execute inference in-memory, avoiding I/O bottlenecks. Database operations utilize "
        "the asynchronous 'motor' driver to log clinician records directly to MongoDB Atlas. Security middleware enforces strict "
        "CORS origin constraints to protect model endpoints."
    )
    
    add_figure(doc, "System_Integration_Workflow.png", "Figure 3: System integration workflow data flow diagram (DFD)")

    # ----------------------------------------------------
    # 6. RESULTS AND RECOMMENDATIONS
    # ----------------------------------------------------
    add_main_heading(doc, "6. Results and Recommendations")
    
    add_body_text(
        doc,
        "Model testing confirms that the proposed deep learning models significantly outperform classical baselines. "
        "The Keras ANN achieved 98.42% accuracy, compared to the Decision Tree (92.30%) and Random Forest (96.50%). "
        "The MobileNetV2 CNN classifier achieved 89.50% validation accuracy. Comparative performance metrics are detailed in Table 1 below:"
    )
    
    # Master Table
    p_tab_cap = doc.add_paragraph()
    p_tab_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tcap = p_tab_cap.add_run("Table 1: Comparative Model Evaluation Metrics")
    run_tcap.font.name = 'Calibri'
    run_tcap.font.size = Pt(10)
    
    table_matrix = doc.add_table(rows=5, cols=6)
    table_matrix.style = 'Light Shading Accent 1'
    hdr = table_matrix.rows[0].cells
    hdr[0].text = 'Model Architecture'
    hdr[1].text = 'Task Target'
    hdr[2].text = 'Accuracy'
    hdr[3].text = 'Precision'
    hdr[4].text = 'Recall'
    hdr[5].text = 'F1-Score'
    
    matrix_rows = [
        ['Decision Tree (Baseline)', 'Symptoms', '92.30%', '92.45%', '92.30%', '92.31%'],
        ['Random Forest (Baseline)', 'Symptoms', '96.50%', '96.60%', '96.50%', '96.52%'],
        ['Artificial Neural Network', 'Symptoms', '98.42%', '98.44%', '98.42%', '98.42%'],
        ['MobileNetV2 CNN', 'Lesion Scans', '89.50%', '88.90%', '89.50%', '88.70%']
    ]
    for row_idx, data in enumerate(matrix_rows):
        row_cells = table_matrix.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text

    add_body_text(
        doc,
        "Visual confusion matrix heatmaps and training history curves are embedded below to verify model performance:"
    )
    
    # Insert matrices & history curves
    add_figure(doc, "Decision_Tree_Confusion_Matrix.png", "Figure 4: Decision Tree confusion matrix heatmap")
    add_figure(doc, "Random_Forest_Confusion_Matrix.png", "Figure 5: Random Forest confusion matrix heatmap")
    add_figure(doc, "ANN_Confusion_Matrix.png", "Figure 6: Symptom Predictor ANN confusion matrix heatmap")
    add_figure(doc, "ANN_Training_History.png", "Figure 7: Symptom Predictor ANN training history curves")
    add_figure(doc, "CNN_Confusion_Matrix.png", "Figure 8: Skin Lesion CNN confusion matrix heatmap")
    add_figure(doc, "CNN_Training_History.png", "Figure 9: Skin Lesion CNN training history curves")

    add_body_text(
        doc,
        "Production verification testing on the deployed system confirms healthy operations. The Vercel frontend has a "
        "load latency of <200ms. The Render backend container memory remains stable at ~340MB, with average model response "
        "latencies of 85ms for the ANN and 120ms for the CNN. Asynchronous log handshakes succeed on MongoDB Atlas."
    )

    # ----------------------------------------------------
    # 7. CONCLUSIONS
    # ----------------------------------------------------
    add_main_heading(doc, "7. Conclusions")
    add_body_text(
        doc,
        "The MediWise AI Capstone Project successfully designs, trains, integrates, and deploys a stateless clinical triage "
        "and lesion scanner portal. Deep learning neural networks are served securely on a python server, ensuring intellectual "
        "property protection and scalable database operations. The platform is fully online and compliant with capstone constraints."
    )

    # ----------------------------------------------------
    # 8. ACKNOWLEDGEMENTS
    # ----------------------------------------------------
    add_main_heading(doc, "8. Acknowledgements")
    add_body_text(
        doc,
        "This project was completed as part of the M.Sc. Software Engineering curriculum. The author acknowledges the support "
        "of the academic advisors and faculty at the University of Europe for Applied Sciences for their guidance throughout "
        "this capstone."
    )

    # ----------------------------------------------------
    # 9. REFERENCES
    # ----------------------------------------------------
    add_main_heading(doc, "9. References")
    references = [
        "1. Keras & TensorFlow APIs. (2025). Keras API Reference. Retrieved from https://keras.io/api/",
        "2. Sandler, M., Howard, A., Zhu, M., Zhmoginov, A., & Chen, L. C. (2018). MobileNetV2: Inverted Residuals and Linear Bottlenecks. CVPR.",
        "3. Tschandl, P., Rosendahl, C., & Kittler, H. (2018). The HAM10000 dataset, a large collection of multi-source dermatoscopic images of common pigmented skin lesions. Scientific Data, 5.",
        "4. FastAPI Python Web Framework. (2026). Documentation. Retrieved from https://fastapi.tiangolo.com/"
    ]
    for ref in references:
        add_body_text(doc, ref)

    output_path = "EBTM881_Capstone_Project_Report.docx"
    doc.save(output_path)
    print(f"Report compiled successfully to {output_path}!")

if __name__ == '__main__':
    build_report()
